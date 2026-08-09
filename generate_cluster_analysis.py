#!/usr/bin/env python3
"""Generate cluster_analysis.docx for Task 3.2 — Thematic Clustering and Comparative Analysis."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# ── Title ───────────────────────────────────────────────────────────────────
title = doc.add_heading('Thematic Clustering and Comparative Analysis', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'This document presents the thematic clustering, cross-cluster comparative analysis, '
    'and future research directions derived from the Literature Analysis Matrix of 50 studies '
    'on computer vision and machine learning models for detecting safety violations in educational institutions.'
)

# ── 1. Table of Thematic Clusters ───────────────────────────────────────────
doc.add_heading('1. Table of Thematic Clusters', level=2)

clusters = [
    {
        'title': 'Deep Learning Architectures',
        'characteristics': 'Studies that design, compare, or optimize neural network topologies (CNN, LSTM, GRU, ResNet, MobileNet, YOLO, 3D CNN, autoencoder) for spatiotemporal feature extraction from surveillance video. Focus on architectural innovation and computational efficiency trade-offs.',
        'representatives': 'Altowairqi et al. (2026); Ul Amin et al. (2024); Dey et al. (2024); Khan et al. (2024); Ul Amin et al. (2022)',
        'findings': 'Hybrid architectures integrating 3D convolutions or pre-trained 2D CNNs with recurrent units (LSTM, BiLSTM, GRU) and attention mechanisms consistently achieve >95% accuracy on controlled benchmarks. Lightweight variants (MobileNet, VD-Net) demonstrate competitive performance with substantially reduced computational requirements suitable for edge deployment.'
    },
    {
        'title': 'Temporal & Motion Analysis',
        'characteristics': 'Studies emphasizing motion representation techniques — optical flow, spatiotemporal feature engineering, action and behavior recognition, and skeleton-based pose estimation — to capture the dynamic dimension of violent or anomalous events.',
        'representatives': 'Mahmoodi & Nezamabadi-Pour (2025); Garcia-Cobo & SanMiguel (2023); Omarov et al. (2022); Park et al. (2024); Yang et al. (2025)',
        'findings': 'Optical flow-based statistical features enable 2D CNNs to achieve performance comparable to more expensive 3D CNNs. Skeleton-based approaches using pose estimation abstract away from appearance-based biases, showing particular promise for child-specific applications where body proportions differ from adult training data.'
    },
    {
        'title': 'Application & Deployment',
        'characteristics': 'Studies focused on real-world deployment of detection systems — violence detection, anomaly detection, weapon/firearm detection — in surveillance contexts (CCTV, IoT, edge computing, crowd analysis) with emphasis on real-time performance, integration with security infrastructure, and practical constraints.',
        'representatives': 'Tapia Leon et al. (2026); Berardini et al. (2024); Vo et al. (2024); Mukto et al. (2024); Abi-Nader et al. (2025)',
        'findings': 'YOLO-family architectures (v5–v8) and SSD-based detectors achieve real-time throughput on commodity and edge hardware (NVIDIA Jetson Nano), with weapon detection mAP scores exceeding 90%. IoT-integrated systems demonstrate successful end-to-end pipelines from detection to automated alerting via Telegram or other messaging platforms.'
    },
    {
        'title': 'Methods & Techniques',
        'characteristics': 'Studies bridging architecture and application through transfer learning, attention mechanisms, feature extraction, pose estimation, and object detection techniques. These works provide methodological foundations that enable the advances reported in the other three clusters.',
        'representatives': 'Shin et al. (2025); Dalal et al. (2024); Singh et al. (2025); Jebur et al. (2023); Aldehim et al. (2023)',
        'findings': 'Attention mechanisms (self-attention, multi-head, channel attention) deliver consistent improvements of 2–4% over non-attention baselines. Transfer learning from ImageNet-pre-trained models significantly reduces data requirements and training time while maintaining high accuracy. Multimodal fusion (RGB + optical flow + audio) outperforms unimodal approaches by approximately 2% average precision.'
    }
]

# Create table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr_cells = table.rows[0].cells
headers = ['Cluster Title', 'Common Characteristics', 'Representative Studies', 'Summary of Key Findings']
for i, text in enumerate(headers):
    hdr_cells[i].text = ''
    p = hdr_cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Header background
    from docx.oxml.ns import qn
    shading = hdr_cells[i]._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): '2B579A',
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)
    run.font.color.rgb = RGBColor(255, 255, 255)

# Data rows
for cluster in clusters:
    row_cells = table.add_row().cells
    row_cells[0].text = ''
    p0 = row_cells[0].paragraphs[0]
    r0 = p0.add_run(cluster['title'])
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.name = 'Times New Roman'

    for j, key in enumerate(['characteristics', 'representatives', 'findings']):
        row_cells[j+1].text = ''
        p = row_cells[j+1].paragraphs[0]
        r = p.add_run(cluster[key])
        r.font.size = Pt(9)
        r.font.name = 'Times New Roman'

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(3.2)
    row.cells[1].width = Cm(4.5)
    row.cells[2].width = Cm(4.0)
    row.cells[3].width = Cm(5.5)

# ── 2. Cross-Cluster Comparative Analysis ───────────────────────────────────
doc.add_heading('2. Cross-Cluster Comparative Analysis', level=2)

comparative = doc.add_paragraph()

p1 = doc.add_paragraph()
p1.add_run(
    'The four thematic clusters exhibit both complementary synergies and meaningful distinctions. '
    'A shared foundation across all clusters is the reliance on deep learning as the core computational '
    'paradigm: whether architecting novel networks (Cluster 1), modeling motion dynamics (Cluster 2), '
    'deploying operational systems (Cluster 3), or advancing methodological techniques (Cluster 4), '
    'every reviewed study depends on convolutional or recurrent neural architectures for feature '
    'representation. This convergence underscores the maturation of deep learning as the de facto '
    'standard for video-based safety violation detection.'
)

p2 = doc.add_paragraph()
p2.add_run(
    'However, notable differences emerge in how the clusters prioritize competing objectives. '
    'Cluster 1 (Architectures) emphasizes accuracy maximization through increasingly sophisticated '
    'network designs, often at the cost of computational complexity. In contrast, Cluster 3 '
    '(Application & Deployment) explicitly privileges real-time throughput and hardware efficiency, '
    'favoring lightweight YOLO and MobileNet variants over the deeper architectures promoted in '
    'Cluster 1. This tension between accuracy and deployability represents a central fault line '
    'in the current literature. Cluster 2 (Temporal & Motion Analysis) and Cluster 4 (Methods & '
    'Techniques) function as methodological bridges: the motion representation strategies from '
    'Cluster 2 provide the temporal understanding that Clusters 1 and 3 depend upon, while the '
    'transfer learning and attention mechanisms from Cluster 4 enhance the performance of models '
    'in all other clusters. Notably, few studies have addressed the intersection of lightweight '
    'architectures (Cluster 1/3) with skeleton-based child-specific modeling (Cluster 2), '
    'representing a significant underexplored area.'
)

p3 = doc.add_paragraph()
p3.add_run(
    'Research gaps identifiable across clusters include: (a) the near-total absence of child-specific '
    'annotated datasets for training and evaluation, as most benchmarks feature adult subjects in '
    'generic public spaces; (b) insufficient attention to privacy-preserving architectures suitable '
    'for monitoring minors in educational settings; (c) limited cross-cluster integration, particularly '
    'combining skeleton-based pose estimation with edge-deployable lightweight models; and (d) a '
    'geographic concentration of research in South and East Asia that constrains the global '
    'generalizability of published findings to diverse educational contexts.'
)

# ── 3. Future Research Recommendations ──────────────────────────────────────
doc.add_heading('3. Future Research Recommendations', level=2)

future = doc.add_paragraph()
future.add_run(
    'Future research should prioritize three directions. First, the creation and open dissemination '
    'of annotated datasets capturing diverse child behaviors in authentic school environments — '
    'building on nascent efforts such as the CABAD and Daily School Break datasets — is essential '
    'for validating model transferability from adult-centric benchmarks to educational settings. '
    'Second, the integration of skeleton-based pose estimation with lightweight edge-deployable '
    'architectures (e.g., quantized YOLO or MobileNet variants) would address the accuracy–efficiency '
    'trade-off that currently fragments the literature across clusters. Third, privacy-preserving '
    'techniques such as on-device processing, federated learning, and differential privacy should '
    'be systematically incorporated into detection pipelines, given that continuous video monitoring '
    'of minors in schools raises distinct ethical and regulatory challenges that the current '
    'literature largely overlooks. Finally, cross-domain robustness evaluations — testing models '
    'across different schools, camera configurations, lighting conditions, and cultural contexts — '
    'should become a standard component of evaluation protocols to ensure that laboratory-validated '
    'systems translate reliably to real-world educational deployments.'
)

# ── Save ────────────────────────────────────────────────────────────────────
doc.save('cluster_analysis.docx')
print("  -> cluster_analysis.docx generated successfully.")
print("✅ Thematic clustering and comparative synthesis successfully generated — ready for integration into the review results chapter.")
