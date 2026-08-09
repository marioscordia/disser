#!/usr/bin/env python3
"""Convert all 5 article section .md files to .docx."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

sections = [
    ('section_introduction.md', 'section_introduction.docx', '1. Introduction'),
    ('section_methodology.md', 'section_methodology.docx', '2. Methodology'),
    ('section_results.md', 'section_results.docx', '3. Results'),
    ('section_discussion.md', 'section_discussion.docx', '4. Discussion'),
    ('section_conclusion.md', 'section_conclusion.docx', '5. Conclusion'),
]

for md_file, docx_file, section_title in sections:
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading(section_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Parse and add paragraphs
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue

        # Skip the heading line if it matches the section title pattern
        if re.match(r'^#\s+\d+\.\s+', line):
            continue

        # Sub-headings (##)
        if line.startswith('## '):
            heading_text = line[3:]
            h = doc.add_heading(heading_text, level=2)
            continue

        # Bold markers: **text**
        # Split on bold markers and add runs
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    doc.save(docx_file)
    print(f"  -> {docx_file}")

print("\nAll 5 sections converted to .docx.")
