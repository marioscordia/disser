#!/usr/bin/env python3
"""Generate literature analysis matrix, top-10 tables, and bibliography."""

import csv
import re
import os
from collections import defaultdict, Counter

# ── Load data ───────────────────────────────────────────────────────────────

def load_papers():
    papers = []
    with open('sourcelist.csv', 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            papers.append(row)
    return papers

def load_references():
    refs = {}
    with open('references.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    entries = re.split(r'\n\n+', content.strip())
    for entry in entries:
        entry = entry.strip()
        if not entry:
            continue
        doi_match = re.search(r'doi\.org/([^\s"]+)', entry)
        if doi_match:
            doi = 'doi.org/' + doi_match.group(1).rstrip('.')
            ref_text = re.sub(r'^\d+\s*', '', entry).strip()
            refs[doi] = ref_text
    return refs

papers = load_papers()
refs = load_references()

# ── Helper: extract info from abstract ──────────────────────────────────────

def extract_aim(abstract):
    """Extract the research aim from abstract."""
    aim_patterns = [
        r'(?:this (?:paper|study|research|article|work)\s*(?:presents|proposes|introduces|aims|focuses|addresses|develops|investigates|explores|evaluates|offers))([^.]*\.)',
        r'(?:we\s*(?:propose|present|introduce|develop|aim|focus|address|investigate|explore|evaluate))([^.]*\.)',
        r'(?:the\s*(?:main\s*)?(?:(?:objective|goal|aim|purpose))[^.]*)([^.]*\.)',
    ]
    for pat in aim_patterns:
        m = re.search(pat, abstract, re.IGNORECASE)
        if m:
            return m.group(0).strip()
    # Fallback: first meaningful sentence
    sentences = abstract.split('.')
    for s in sentences[:2]:
        s = s.strip()
        if len(s) > 40:
            return s + '.'
    return 'Not specified'

def extract_method(abstract):
    """Extract methodology from abstract."""
    method_keywords = [
        r'(CNN|Convolutional Neural Network|convolutional neural network)',
        r'(LSTM|Long Short-Term Memory|long short-term memory)',
        r'(YOLO|you only look once)',
        r'(ResNet|Residual Network)',
        r'(autoencoder|auto-encoder)',
        r'(transformer|attention mechanism)',
        r'(transfer learning)',
        r'(3D\s*(?:CNN|convolutional))',
        r'(GRU|Gated Recurrent Unit)',
        r'(MobileNet)',
        r'(Inception)',
        r'(optical flow)',
        r'(pose estimation|human pose)',
        r'(BiLSTM|bidirectional LSTM)',
        r'(DenseNet)',
        r'(SVM|support vector)',
        r'(GAN|generative adversarial)',
    ]
    found = []
    for pat in method_keywords:
        if re.search(pat, abstract, re.IGNORECASE):
            kw = re.search(pat, abstract, re.IGNORECASE).group(0)
            if kw not in found:
                found.append(kw)
    if found:
        return 'Uses ' + ', '.join(found[:5]) + '.'
    return 'Deep learning-based approach.'

def extract_dataset(abstract, title=''):
    """Extract dataset/participants information."""
    dataset_patterns = [
        r'(?:UCF-Crime|UCF Crime)',
        r'ShanghaiTech(?:\s*Campus)?',
        r'(?:Violent Flows|ViolentFlow)',
        r'(?:Hockey\s*(?:Fight|Fights))',
        r'(?:Movie\s*Fights|MovieFights)',
        r'XD-Violence',
        r'UCSD\s*(?:Ped[12]|PED[12])',
        r'CUHK\s*Avenue',
        r'(?:RWF-2000|RWF2000)',
        r'RLVS',
        r'UBI-Fight',
        r'(?:Crowd|CrowdSurge)',
        r'IITR-IAR',
        r'AIRTLab',
        r'SCVD',
        r'EAVDD',
        r'DSB',
        r'CABAD',
    ]
    text = title + ' ' + abstract
    found = []
    for pat in dataset_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            found.append(m.group(0))
    if found:
        return ', '.join(list(set(found))[:4])
    # Check for custom dataset mentions
    if re.search(r'custom\s*(?:dataset|weapon)', text, re.IGNORECASE):
        return 'Custom dataset'
    if re.search(r'(?:\d{1,3}(?:,\d{3})*\s*(?:images|videos|clips))', text):
        m = re.search(r'(?:\d{1,3}(?:,\d{3})*\s*(?:images|videos|clips))', text, re.IGNORECASE)
        return m.group(0)
    return 'Not specified'

def extract_findings(abstract):
    """Extract main findings from abstract."""
    findings_patterns = [
        r'(?:results\s*(?:demonstrate|show|indicate|reveal|highlight|confirm|validate|prove|suggest))([^.]*\d{1,3}(?:\.\d)?%?[^.]*\.)',
        r'(?:achiev(?:es|ing|ed)\s*(?:an?\s*)?(?:accuracy|AUC|F1|AP|mAP|precision)[^.]*\.)',
        r'(?:outperforms?[^.]*\.)',
        r'(?:achieves?\s*(?:\d{1,3}(?:\.\d)?%?\s*(?:accuracy|AUC|F1|AP|mAP|precision|detection|recall))[^.]*\.)',
    ]
    for pat in findings_patterns:
        m = re.search(pat, abstract, re.IGNORECASE)
        if m:
            return m.group(0).strip()
    # Fallback: last meaningful sentence with numbers
    sentences = abstract.split('.')
    for s in reversed(sentences):
        s = s.strip()
        if re.search(r'\d{1,3}(?:\.\d)?%', s) and len(s) > 30:
            return s + '.'
    return 'See paper for details.'

def extract_limitations(abstract):
    """Identify limitations from abstract."""
    limit_patterns = [
        r'(?:limitation|limited|challenge|challenging|struggle|computational\s*(?:complexity|cost)|lack|scarce|scarcity)[^.]*\.',
        r'(?:future\s*(?:research|work|studies|directions?))[^.]*\.',
    ]
    for pat in limit_patterns:
        m = re.search(pat, abstract, re.IGNORECASE)
        if m:
            return m.group(0).strip()
    return 'Not specified'

def extract_keywords_from_text(title, abstract):
    """Extract actual keywords from title and abstract text."""
    curated = [
        'violence detection', 'anomaly detection', 'weapon detection', 'firearm detection',
        'deep learning', 'CNN', 'LSTM', 'YOLO', 'convolutional neural network',
        'video surveillance', 'object detection', 'optical flow', 'attention mechanism',
        'transfer learning', 'real-time detection', 'pose estimation', 'GRU', 'autoencoder',
        'spatiotemporal features', 'behavior recognition', 'action recognition',
        'edge computing', 'IoT', 'CCTV', 'crowd analysis', 'skeleton-based', 'ResNet',
        'MobileNet', '3D CNN', 'feature extraction'
    ]
    text = (title + ' ' + abstract).lower()
    found = []
    for kw in curated:
        if kw.lower() in text and len(found) < 6:
            found.append(kw)
    return '; '.join(found[:5]) if found else 'deep learning; computer vision'

# ═════════════════════════════════════════════════════════════════════════════
# TASK 3.1: Literature Analysis Matrix
# ═════════════════════════════════════════════════════════════════════════════

def generate_matrix():
    rows = []
    for p in papers:
        title = (p.get('Title', '') or '').strip()
        abstract = (p.get('Short description', '') or '').strip()
        doi = (p.get('DOI', '') or '').strip()
        authors = (p.get('Authors', '') or '').strip()
        year = (p.get('Year', '') or '').strip()

        # Match APA reference
        apa = ''
        for ref_doi, ref_text in refs.items():
            if doi and doi.lower() in ref_doi.lower():
                apa = ref_text
                break
            if doi and ref_doi.lower() in doi.lower():
                apa = ref_text
                break

        # If no DOI match, try title match in references
        if not apa:
            title_words = title[:60].lower()
            for ref_doi, ref_text in refs.items():
                if title_words[:40] in ref_text.lower().replace('<i>','').replace('</i>','').replace('<scp>','').replace('</scp>',''):
                    apa = ref_text
                    break

        row = {
            'Author(s)': authors,
            'Year': year,
            'Aim': extract_aim(abstract) if abstract else 'Not specified',
            'Method': extract_method(abstract) if abstract else 'Not specified',
            'Participants/Dataset': extract_dataset(abstract, title),
            'Main Findings': extract_findings(abstract) if abstract else 'Not specified',
            'Key Limitations': extract_limitations(abstract) if abstract else 'Not specified',
            'Keywords': extract_keywords_from_text(title, abstract),
            'APA Reference': apa
        }
        rows.append(row)

    # Write CSV
    fieldnames = ['Author(s)', 'Year', 'Aim', 'Method', 'Participants/Dataset',
                   'Main Findings', 'Key Limitations', 'Keywords', 'APA Reference']
    with open('matrix_literature.csv', 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)

    print(f"  -> matrix_literature.csv ({len(rows)} rows)")

generate_matrix()

# ═════════════════════════════════════════════════════════════════════════════
# TASK 7.1: Top 10 Authors
# ═════════════════════════════════════════════════════════════════════════════

def generate_top_authors():
    author_papers = defaultdict(list)
    for p in papers:
        authors_str = (p.get('Authors', '') or '').strip()
        if not authors_str:
            continue
        title = (p.get('Title', '') or '').strip()
        year = int(p.get('Year', 2024) or 2024)
        authors = [a.strip() for a in authors_str.split(';') if a.strip()]
        for a in authors:
            author_papers[a].append((title, year))

    # Sort by paper count
    sorted_authors = sorted(author_papers.items(), key=lambda x: len(x[1]), reverse=True)

    rows = []
    for author, papers_list in sorted_authors[:10]:
        paper_count = len(papers_list)
        # Most recent paper
        recent = max(papers_list, key=lambda x: x[1])
        # Research focus from titles
        titles = [t for t, _ in papers_list]
        focus = infer_focus(titles)
        # Contribution summary
        contrib = infer_contribution(author, titles)

        rows.append({
            'Author': author,
            'Paper Count': paper_count,
            'Research Focus': focus,
            'Key Paper': recent[0],
            'Contribution Summary': contrib
        })

    with open('top10_authors.csv', 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=['Author', 'Paper Count', 'Research Focus', 'Key Paper', 'Contribution Summary'])
        writer.writeheader()
        writer.writerows(rows)

    print(f"  -> top10_authors.csv ({len(rows)} rows)")

def infer_focus(titles):
    """Infer research focus from paper titles."""
    text = ' '.join(titles).lower()
    focuses = []
    if 'violence' in text or 'violent' in text: focuses.append('violence detection')
    if 'anomaly' in text: focuses.append('anomaly detection')
    if 'weapon' in text or 'firearm' in text or 'gun' in text: focuses.append('weapon detection')
    if 'behavior' in text or 'behaviour' in text: focuses.append('behavior analysis')
    if 'surveillance' in text: focuses.append('video surveillance')
    return '; '.join(focuses) if focuses else 'video analysis'

def infer_contribution(author, titles):
    """Generate a 1-sentence contribution summary."""
    focus = infer_focus(titles)
    count = len(titles)
    return f"Contributed {count} paper(s) on {focus}, advancing deep learning and computer vision applications for automated surveillance and public safety."

generate_top_authors()

# ═════════════════════════════════════════════════════════════════════════════
# TASK 7.2: Top 10 Keywords
# ═════════════════════════════════════════════════════════════════════════════

def generate_top_keywords():
    keyword_counts = Counter()
    for p in papers:
        title = (p.get('Title', '') or '').lower()
        abstract = (p.get('Short description', '') or '').lower()
        text = title + ' ' + abstract
        found = set()
        curated = CURATED_KEYWORDS
        for kw in curated:
            if kw.lower() in text:
                found.add(kw)
        for kw in found:
            keyword_counts[kw] += 1

    definitions = {
        'CNN': 'A class of deep neural networks most commonly applied to analyze visual imagery, using convolutional layers to automatically learn spatial hierarchies of features.',
        'LSTM': 'A type of recurrent neural network architecture capable of learning long-term dependencies in sequential data, commonly used for temporal modeling in video analysis.',
        'YOLO': 'A real-time object detection system that frames object detection as a single regression problem, predicting bounding boxes and class probabilities directly from full images.',
        'deep learning': 'A subset of machine learning based on artificial neural networks with multiple layers, enabling automatic feature extraction from raw data.',
        'video surveillance': 'The use of cameras to monitor activities in public or private spaces, increasingly augmented by AI-driven automated analysis.',
        'anomaly detection': 'The identification of rare events, items, or observations that differ significantly from the majority of data, critical for security applications.',
        'violence detection': 'The automated identification of aggressive physical behavior in video streams using computer vision and machine learning techniques.',
        'object detection': 'A computer vision task involving both localization and classification of objects within an image or video frame.',
        'attention mechanism': 'A technique that enables neural networks to focus on the most relevant parts of the input data, improving feature representation and model interpretability.',
        'optical flow': 'The pattern of apparent motion of objects, surfaces, and edges in a visual scene caused by relative motion between an observer and the scene.',
        'transfer learning': 'A machine learning method where a model developed for one task is reused as the starting point for a model on a second task, reducing training data requirements.',
        'feature extraction': 'The process of transforming raw data into a set of characteristic representations suitable for machine learning models.',
        'real-time detection': 'The ability of a system to process video frames and produce detection results with minimal latency, suitable for live surveillance.',
        'convolutional neural network': 'Same as CNN.',
        'pose estimation': 'Computer vision technique to determine the position and orientation of the human body and its parts from images or video.',
        'spatiotemporal features': 'Data representations that capture both spatial (appearance) and temporal (motion) information from video sequences.',
        'action recognition': 'The task of identifying specific human actions or activities from video data.',
        'autoencoder': 'An unsupervised neural network architecture used for learning efficient data encodings, often applied in anomaly detection through reconstruction error.',
        'weapon detection': 'Automated identification of weapons, particularly firearms and knives, in surveillance footage using object detection models.',
        'firearm detection': 'A specific subset of weapon detection focused on identifying guns and firearms in images or video streams.',
        'IoT': 'A network of physical objects embedded with sensors and software to connect and exchange data, enabling distributed surveillance systems.',
        'edge computing': 'Distributed computing paradigm that brings computation and data storage closer to the data source, reducing latency for real-time surveillance.',
        'CCTV': 'Closed-circuit television systems used for surveillance monitoring; increasingly integrated with AI-based automated analysis.',
        'ResNet': 'Deep residual network architecture using skip connections to enable training of very deep networks, widely used as a backbone for vision tasks.',
        'MobileNet': 'Efficient CNN architecture designed for mobile and embedded vision applications, balancing accuracy with computational cost.',
        '3D CNN': 'Extension of 2D convolutional networks to process volumetric data or video by applying 3D filters that capture both spatial and temporal dimensions.',
        'GRU': 'A gating mechanism in recurrent neural networks, similar to LSTM but with a simpler architecture, used for sequential data processing.',
        'behavior recognition': 'The automated classification of human behavior patterns from video, including both normal activities and suspicious or anomalous behaviors.',
        'skeleton-based': 'Methods that use human body skeleton representations derived from pose estimation for activity and behavior recognition.',
        'crowd analysis': 'The study of crowd behavior using computer vision techniques, including density estimation, counting, and anomaly detection in crowded scenes.',
    }

    category_map = {
        'CNN': 'Architecture', 'LSTM': 'Architecture', 'YOLO': 'Architecture',
        'GRU': 'Architecture', 'ResNet': 'Architecture', 'MobileNet': 'Architecture',
        '3D CNN': 'Architecture', 'autoencoder': 'Architecture', 'convolutional neural network': 'Architecture',
        'deep learning': 'Method', 'attention mechanism': 'Method', 'transfer learning': 'Method',
        'pose estimation': 'Method', 'feature extraction': 'Method', 'object detection': 'Method',
        'optical flow': 'Method', 'spatiotemporal features': 'Method', 'skeleton-based': 'Method',
        'violence detection': 'Application', 'anomaly detection': 'Application',
        'weapon detection': 'Application', 'firearm detection': 'Application',
        'behavior recognition': 'Application', 'action recognition': 'Application',
        'crowd analysis': 'Application',
        'video surveillance': 'Context', 'CCTV': 'Context', 'real-time detection': 'Context',
        'IoT': 'Context', 'edge computing': 'Context',
    }

    top10 = keyword_counts.most_common(10)
    rows = []
    for kw, freq in top10:
        rows.append({
            'Keyword': kw,
            'Frequency': freq,
            'Definition': definitions.get(kw, 'A key term in computer vision and surveillance research.'),
            'Relevance to Topic': f'Highly relevant to automated surveillance and safety violation detection using computer vision.',
            'Category': category_map.get(kw, 'Method')
        })

    with open('top10_keywords.csv', 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=['Keyword', 'Frequency', 'Definition', 'Relevance to Topic', 'Category'])
        writer.writeheader()
        writer.writerows(rows)

    print(f"  -> top10_keywords.csv ({len(rows)} rows)")

CURATED_KEYWORDS = [
    'violence detection', 'anomaly detection', 'weapon detection', 'firearm detection',
    'deep learning', 'CNN', 'LSTM', 'YOLO', 'convolutional neural network',
    'video surveillance', 'object detection', 'optical flow', 'attention mechanism',
    'transfer learning', 'real-time detection', 'pose estimation', 'GRU', 'autoencoder',
    'spatiotemporal features', 'behavior recognition', 'action recognition',
    'edge computing', 'IoT', 'CCTV', 'crowd analysis', 'skeleton-based', 'ResNet',
    'MobileNet', '3D CNN', 'feature extraction'
]

generate_top_keywords()

# ═════════════════════════════════════════════════════════════════════════════
# TASK 2.4: Bibliography
# ═════════════════════════════════════════════════════════════════════════════

def generate_bibliography():
    ref_text = ''
    with open('references.txt', 'r', encoding='utf-8') as f:
        ref_text = f.read()

    # Split into individual references
    entries = re.split(r'\n\n+', ref_text.strip())

    lines = ['# Bibliography\n', '> APA-formatted reference list for all 50 papers included in the systematic review.\n']

    for i, entry in enumerate(entries):
        entry = entry.strip()
        if not entry:
            continue
        # Remove leading number
        clean = re.sub(r'^\d+\s*', '', entry).strip()
        # Clean up HTML/XML tags from API
        clean = clean.replace('<i>', '').replace('</i>', '')
        clean = clean.replace('<scp>', '').replace('</scp>', '')
        clean = re.sub(r'\s+', ' ', clean).strip()
        lines.append(f'{i+1}. {clean}\n')

    with open('bibliography.md', 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    # Also generate .docx if python-docx available
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Bibliography', level=1)
        doc.add_paragraph('APA-formatted reference list for all 50 papers included in the systematic review.')
        for i, entry in enumerate(entries):
            entry = entry.strip()
            if not entry:
                continue
            clean = re.sub(r'^\d+\s*', '', entry).strip()
            clean = clean.replace('<i>', '').replace('</i>', '')
            clean = clean.replace('<scp>', '').replace('</scp>', '')
            clean = re.sub(r'\s+', ' ', clean).strip()
            doc.add_paragraph(f'{i+1}. {clean}')
        doc.save('bibliography.docx')
        print("  -> bibliography.docx")
    except Exception:
        pass

    print("  -> bibliography.md")

generate_bibliography()

print("\nAll data files generated successfully.")
